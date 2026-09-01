import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>\n'
                          f'<w:top w:val="{kwargs.get("top", "none")}" w:sz="{kwargs.get("top_sz", "4")}" w:space="0" w:color="{kwargs.get("top_color", "auto")}"/>\n'
                          f'<w:left w:val="{kwargs.get("left", "none")}" w:sz="{kwargs.get("left_sz", "4")}" w:space="0" w:color="{kwargs.get("left_color", "auto")}"/>\n'
                          f'<w:bottom w:val="{kwargs.get("bottom", "none")}" w:sz="{kwargs.get("bottom_sz", "4")}" w:space="0" w:color="{kwargs.get("bottom_color", "auto")}"/>\n'
                          f'<w:right w:val="{kwargs.get("right", "none")}" w:sz="{kwargs.get("right_sz", "4")}" w:space="0" w:color="{kwargs.get("right_color", "auto")}"/>\n'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)

def update_toc_in_docx(file_path):
    doc = docx.Document(file_path)
    
    # 1. Clean out existing TOC elements if already present
    # Find paragraph "Diagram สถาปัตยกรรมของระบบ"
    target_idx = None
    for idx, p in enumerate(doc.paragraphs):
        if 'Diagram สถาปัตยกรรมของระบบ' in p.text:
            target_idx = idx
            break
            
    if target_idx is None:
        print(f"Could not find target paragraph in {file_path}")
        return

    # Delete any paragraphs between cover page (after '31 สิงหาคม 2569') and target_idx
    cover_idx = None
    for idx, p in enumerate(doc.paragraphs):
        if '31 สิงหาคม 2569' in p.text:
            cover_idx = idx
            break
            
    if cover_idx is not None and target_idx > cover_idx + 1:
        # Delete elements between cover_idx and target_idx
        p_to_remove = doc.paragraphs[cover_idx+1:target_idx]
        for p in p_to_remove:
            p._element.getparent().remove(p._element)
            
    # Also remove the first 3 tables if they are TOC tables
    # Check if table 0 is TOC (has 2 cols and rows >= 20)
    while len(doc.tables) > 0 and len(doc.tables[0].columns) == 2 and len(doc.tables[0].rows) in [34, 21, 6]:
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

    # Re-find target paragraph
    target_p = None
    for p in doc.paragraphs:
        if 'Diagram สถาปัตยกรรมของระบบ' in p.text:
            target_p = p
            break

    # Helpers
    def insert_p(text='', bold=False, size=16, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
        new_p = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
        if align != WD_ALIGN_PARAGRAPH.LEFT:
            new_p.alignment = align
        pPr = new_p._p.get_or_add_pPr()
        spPr = parse_xml(f'<w:spacing {nsdecls("w")} w:before="{int(space_before*20)}" w:after="{int(space_after*20)}"/>')
        pPr.append(spPr)
        
        if text:
            run = new_p.add_run(text)
            run.bold = bold
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(size)
            rPr = run._r.get_or_add_rPr()
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="TH Sarabun New" w:hAnsi="TH Sarabun New" w:cs="TH Sarabun New" w:eastAsia="TH Sarabun New"/>')
            rPr.append(rFonts)
            
        target_p._p.addprevious(new_p._p)
        return new_p

    def insert_page_break():
        new_p = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
        run = new_p.add_run()
        run.add_break(docx.enum.text.WD_BREAK.PAGE)
        target_p._p.addprevious(new_p._p)

    def insert_table(rows_data, col_widths, headers=('หัวข้อ', 'หน้า')):
        new_tbl = doc.add_table(rows=len(rows_data)+1, cols=len(col_widths))
        new_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        hdr_cells = new_tbl.rows[0].cells
        for idx, h_text in enumerate(headers):
            hdr_cells[idx].text = h_text
            p = hdr_cells[idx].paragraphs[0]
            if idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.runs[0].bold = True
            p.runs[0].font.name = 'TH Sarabun New'
            p.runs[0].font.size = Pt(15)
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="TH Sarabun New" w:hAnsi="TH Sarabun New" w:cs="TH Sarabun New" w:eastAsia="TH Sarabun New"/>')
            p.runs[0]._r.get_or_add_rPr().append(rFonts)
            set_cell_border(hdr_cells[idx], bottom='single', bottom_sz='6')
            
        # Data rows
        for r_idx, (title, page_num, is_bold, indent_level) in enumerate(rows_data):
            row_cells = new_tbl.rows[r_idx+1].cells
            
            # Col 0 (Title)
            p0 = row_cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(1)
            p0.paragraph_format.space_after = Pt(1)
            indent_str = '    ' * indent_level
            run0 = p0.add_run(indent_str + title)
            run0.bold = is_bold
            run0.font.name = 'TH Sarabun New'
            run0.font.size = Pt(14 if not is_bold else 15)
            rFonts0 = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="TH Sarabun New" w:hAnsi="TH Sarabun New" w:cs="TH Sarabun New" w:eastAsia="TH Sarabun New"/>')
            run0._r.get_or_add_rPr().append(rFonts0)
            
            # Col 1 (Page)
            p1 = row_cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p1.paragraph_format.space_before = Pt(1)
            p1.paragraph_format.space_after = Pt(1)
            run1 = p1.add_run(str(page_num))
            run1.bold = is_bold
            run1.font.name = 'TH Sarabun New'
            run1.font.size = Pt(14 if not is_bold else 15)
            rFonts1 = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="TH Sarabun New" w:hAnsi="TH Sarabun New" w:cs="TH Sarabun New" w:eastAsia="TH Sarabun New"/>')
            run1._r.get_or_add_rPr().append(rFonts1)
            
            set_cell_border(row_cells[0], bottom='none')
            set_cell_border(row_cells[1], bottom='none')

        for row in new_tbl.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

        target_p._p.addprevious(new_tbl._tbl)

    # 1. Page break after cover page
    insert_page_break()

    # 2. Main TOC (Starts at Page 5)
    insert_p('สารบัญ', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    
    toc_data = [
        ('1. Diagram สถาปัตยกรรมของระบบ', '5', True, 0),
        ('1.1 ภาพรวมกระบวนการหลักของระบบ', '5', False, 1),
        ('1.2 ภาพรวมของระบบ (System Overview)', '5', False, 1),
        ('1.3 Data Flow การอ่าน (Read Path)', '6', False, 1),
        ('1.4 Data Flow การเขียน (Write Path)', '6', False, 1),
        ('1.5 เทคโนโลยีและ Dependencies ที่ใช้ในระบบ', '7', False, 1),
        ('2. กลยุทธ์ Cache Invalidation และการป้องกันการสั่งซื้อซ้ำซ้อน', '8', True, 0),
        ('2.1 ปัญหา remainingStock เปลี่ยนเร็วมากแต่ถูกอ่านหนักมาก', '8', False, 1),
        ('2.2 วิธีแก้: แยกข้อมูลตามอัตราการเปลี่ยนแปลง (Cache-Aside + Stock Overlay)', '8', False, 1),
        ('2.3 กลยุทธ์ Cache Invalidation', '10', False, 1),
        ('2.3.1 กลไกล้างแคชด้วย Live-Key Index (ไม่ใช้คำสั่ง KEYS)', '11', False, 2),
        ('2.3.2 การป้องกัน Cache Stampede (Single-Flight & TTL Jitter)', '12', False, 2),
        ('2.4 การป้องกันการสั่งซื้อซ้ำซ้อนและการขายเกิน (4-Tier Defense)', '12', False, 1),
        ('2.4.1 Tier 1: Atomic Lua Gatekeeper (ชั้นที่ตัดสินจริง)', '13', False, 2),
        ('2.4.2 กลไกป้องกันการซื้อซ้ำ', '14', False, 2),
        ('2.4.3 Tier 2: BullMQ (Asynchronous Queue & Idempotency)', '14', False, 2),
        ('2.4.4 Tier 3: Worker & Atomic SQL Decrement (WHERE remaining_stock > 0)', '15', False, 2),
        ('2.4.5 Tier 4: Database Constraints (UNIQUE & CHECK Constraints)', '16', False, 2),
        ('3. ผลลัพธ์จาก Load Test Dashboard', '17', True, 0),
        ('3.1 สรุปผลตัวชี้วัดเทียบเกณฑ์ (Local 6 instances)', '17', False, 1),
        ('3.2 k6 Summary Dashboard', '18', False, 1),
        ('3.3 Bull-Board สถานะคิวงาน', '18', False, 1),
        ('3.4 สถิติ Cache Hit / Miss', '19', False, 1),
        ('3.5 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 01', '19', False, 1),
        ('3.6 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 02', '20', False, 1),
        ('3.7 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 03', '21', False, 1),
        ('3.8 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 04', '22', False, 1),
        ('3.9 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 05', '23', False, 1),
        ('3.10 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 07', '24', False, 1),
        ('3.11 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 08', '25', False, 1),
        ('3.12 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 09', '26', False, 1),
        ('3.13 Observability Dashboard (/admin/queues, /admin/insights)', '27', False, 1),
        ('4. สรุปผลการทดสอบและข้อเสนอแนะ', '28', True, 0)
    ]
    insert_table(toc_data, [5.5, 1.0], ('ชื่อเรื่อง', 'หน้า'))

    # 3. List of Figures
    insert_page_break()
    insert_p('สารบัญภาพ', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    
    figures_data = [
        ('ภาพที่ 1 ภาพรวมกระบวนการหลักของระบบ', '5', False, 0),
        ('ภาพที่ 2 สถาปัตยกรรมภาพรวมของระบบ Flash Sale', '5', False, 0),
        ('ภาพที่ 3 Data Flow Diagram เส้นทางการอ่าน (Read Path)', '6', False, 0),
        ('ภาพที่ 4 Data Flow Diagram เส้นทางการเขียน (Write Path)', '6', False, 0),
        ('ภาพที่ 5 ขั้นตอนการอ่าน Cache-Aside + Stock Overlay', '8', False, 0),
        ('ภาพที่ 6 สถาปัตยกรรมป้องกันการขายเกิน 4-Tier Defense', '12', False, 0),
        ('ภาพที่ 7 ผลสรุปการทดสอบจาก k6 Summary Dashboard', '18', False, 0),
        ('ภาพที่ 8 สถานะคิวงานจาก Bull-Board (Completed = 50)', '18', False, 0),
        ('ภาพที่ 9 สถิติ Cache Hit / Miss', '19', False, 0),
        ('ภาพที่ 10 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 01', '19', False, 0),
        ('ภาพที่ 11 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 02', '20', False, 0),
        ('ภาพที่ 12 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 03', '21', False, 0),
        ('ภาพที่ 13 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 04', '22', False, 0),
        ('ภาพที่ 14 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 05', '23', False, 0),
        ('ภาพที่ 15 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 07', '24', False, 0),
        ('ภาพที่ 16 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 08', '25', False, 0),
        ('ภาพที่ 17 ผลการยิงข้ามเครือข่ายไปยัง Cloud VM กลุ่ม 09', '26', False, 0),
        ('ภาพที่ 18 หน้า Dashboard path /admin/queues (Bull-Board)', '27', False, 0),
        ('ภาพที่ 19 หน้า Dashboard path /admin/insights — สถานะสต็อกและ drift (1)', '27', False, 0),
        ('ภาพที่ 20 หน้า Dashboard path /admin/insights — Connection Pool & Replication (2)', '28', False, 0)
    ]
    insert_table(figures_data, [5.5, 1.0], ('ภาพที่', 'หน้า'))

    # 4. List of Tables
    insert_page_break()
    insert_p('สารบัญตาราง', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    
    tables_data = [
        ('ตารางที่ 1 การแยกชนิดข้อมูลตามอัตราการเปลี่ยนแปลง (Cache-Aside vs Stock Overlay)', '8', False, 0),
        ('ตารางที่ 2 นโยบายการ Invalidate แคชตามประเภทของข้อมูล', '10', False, 0),
        ('ตารางที่ 3 กลไกและระยะเวลา TTL ในการป้องกันการซื้อซ้ำซ้อน', '14', False, 0),
        ('ตารางที่ 4 สรุปผลตัวชี้วัดเทียบเกณฑ์มาตรฐาน (Benchmark Results)', '17', False, 0),
        ('ตารางที่ 5 รายละเอียดหน้า Observability Dashboard ที่มีในระบบ', '27', False, 0)
    ]
    insert_table(tables_data, [5.5, 1.0], ('ตารางที่', 'หน้า'))

    # Final page break before Section 1
    insert_page_break()

    doc.save(file_path)
    print(f'Successfully updated TOC (starts at page 5) in {file_path}')

if __name__ == '__main__':
    update_toc_in_docx(r'docs/Report/สาขาวิชาวิศวกรรมคอมพิวเตอร์.docx')
    update_toc_in_docx(r'C:\Users\Acer\Downloads\สาขาวิชาวิศวกรรมคอมพิวเตอร์.docx')
